"""DOCX reader. Yields one Chunk per paragraph and one per table cell.

Table cells include their column header (the first row of the table) so the
header heuristic can fire. python-docx ignores headers/footers and tracked
changes — those are out of scope.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterator

from docx import Document

from ..models import Chunk


def read(path: Path) -> Iterator[Chunk]:
    doc = Document(str(path))

    for para_num, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        yield Chunk(
            file_path=str(path),
            file_type="docx",
            location=f"paragraph={para_num}",
            text=text,
        )

    for tbl_idx, table in enumerate(doc.tables, start=1):
        rows = list(table.rows)
        if not rows:
            continue
        header_cells = [c.text.strip() for c in rows[0].cells]
        for row_num, row in enumerate(rows[1:], start=2):
            for col_idx, cell in enumerate(row.cells):
                value = cell.text.strip()
                if not value:
                    continue
                header = header_cells[col_idx] if col_idx < len(header_cells) else None
                yield Chunk(
                    file_path=str(path),
                    file_type="docx",
                    location=f"table={tbl_idx} row={row_num} col={header or col_idx + 1}",
                    text=value,
                    column_header=header,
                )
