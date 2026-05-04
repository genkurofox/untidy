"""Tests for the new readers: JSON/NDJSON/YAML, PDF, DOCX."""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from untidy.readers import json_reader
from untidy.scanner import scan


# --- JSON / NDJSON / YAML --------------------------------------------------

def test_json_reader_uses_keys_as_headers(tmp_path: Path):
    p = tmp_path / "patient.json"
    p.write_text(json.dumps({
        "first_name": "Jane",
        "last_name": "Smith",
        "ssn": "123-45-6789",
        "address": {"street": "456 Mission Ave", "zip": "94105"},
    }))
    chunks = list(json_reader.read(p))
    headers_seen = {c.column_header for c in chunks}
    assert "first_name" in headers_seen
    assert "ssn" in headers_seen
    assert "street" in headers_seen
    # Path-style locations
    assert any(c.location == "first_name" for c in chunks)
    assert any(c.location.startswith("address.") for c in chunks)


def test_json_scan_finds_pii_via_key_header(tmp_path: Path):
    p = tmp_path / "rec.json"
    p.write_text(json.dumps({"first_name": "Jane", "ssn": "123-45-6789"}))
    findings = list(scan([p]))
    types = {f.entity_type for f in findings}
    assert "NAME_FROM_HEADER" in types
    assert "SSN" in types


def test_ndjson_scan(tmp_path: Path):
    p = tmp_path / "events.ndjson"
    p.write_text(
        json.dumps({"email": "a@b.com", "ssn": "123-45-6789"}) + "\n"
        + json.dumps({"email": "c@d.com"}) + "\n"
    )
    findings = list(scan([p]))
    emails = [f for f in findings if f.entity_type == "EMAIL"]
    assert len(emails) == 2


def test_yaml_scan(tmp_path: Path):
    pytest.importorskip("yaml")
    p = tmp_path / "config.yaml"
    p.write_text(
        "credentials:\n"
        "  api_key: AKIAIOSFODNN7EXAMPLE\n"
        "  user_email: admin@example.com\n"
    )
    findings = list(scan([p]))
    types = {f.entity_type for f in findings}
    assert "AWS_ACCESS_KEY_ID" in types
    assert "EMAIL" in types


# --- PDF -------------------------------------------------------------------

def test_pdf_scan(tmp_path: Path):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    p = tmp_path / "letter.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 720, "Patient SSN: 123-45-6789")
    c.drawString(72, 700, "Email: jane@example.com")
    c.showPage()
    c.save()

    findings = list(scan([p]))
    types = {f.entity_type for f in findings}
    assert "SSN" in types
    assert "EMAIL" in types
    assert all(f.location.startswith("page=") for f in findings)


# --- DOCX ------------------------------------------------------------------

def test_docx_scan_paragraphs_and_tables(tmp_path: Path):
    from docx import Document

    p = tmp_path / "intake.docx"
    doc = Document()
    doc.add_paragraph("Patient name: Maria Garcia")
    doc.add_paragraph("DOB: 1988-02-19")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "ssn"
    table.rows[0].cells[1].text = "email"
    table.rows[1].cells[0].text = "456-78-9012"
    table.rows[1].cells[1].text = "maria@example.com"
    doc.save(str(p))

    findings = list(scan([p]))
    types = {f.entity_type for f in findings}
    assert "NAME" in types  # from inline NAME pattern
    assert "DOB" in types
    assert "SSN" in types
    assert "EMAIL" in types
