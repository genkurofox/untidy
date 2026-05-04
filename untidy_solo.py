#!/usr/bin/env python3
"""untidy_solo — single-file offline scanner for sensitive data.

A drop-in version of the `untidy` package: same CLI, same detection logic,
zero install required for the core file types. Optional dependencies unlock
extra file types and degrade gracefully if missing:

  CSV / SQL / JSON / NDJSON / text / log / md  — stdlib only
  Excel (.xlsx)                                — needs openpyxl
  PDF (.pdf)                                   — needs pdfminer.six
  DOCX (.docx)                                 — needs python-docx
  YAML (.yaml/.yml)                            — needs PyYAML

Usage:
    python untidy_solo.py scan PATH [PATH ...] [options]
    python untidy_solo.py scan-git REPO_PATH [options]

Exit codes: 0 = clean, 1 = findings, 2 = error (or read-error with --strict).
"""
from __future__ import annotations

import argparse
import csv as _csv
import fnmatch
import json
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable, Iterator, Optional

__version__ = "0.1.0-solo"


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class Chunk:
    file_path: str
    file_type: str
    location: str
    text: str
    column_header: Optional[str] = None


@dataclass(frozen=True)
class Finding:
    file_path: str
    file_type: str
    location: str
    entity_type: str
    detection_rule: str
    confidence: str
    match_snippet: str
    column_header: Optional[str] = None


# ---------------------------------------------------------------------------
# Checksums
# ---------------------------------------------------------------------------

def luhn_valid(s: str) -> bool:
    d = [int(c) for c in s if c.isdigit()]
    if len(d) < 13 or len(d) > 19:
        return False
    total = 0
    for i, n in enumerate(reversed(d)):
        if i % 2 == 1:
            n *= 2
            if n > 9:
                n -= 9
        total += n
    return total % 10 == 0


def ssn_valid(s: str) -> bool:
    digits = "".join(c for c in s if c.isdigit())
    if len(digits) != 9:
        return False
    area, group, serial = digits[:3], digits[3:5], digits[5:]
    if area in {"000", "666"} or area.startswith("9"):
        return False
    if group == "00" or serial == "0000":
        return False
    return True


def routing_valid(s: str) -> bool:
    d = [int(c) for c in s if c.isdigit()]
    if len(d) != 9:
        return False
    checksum = (
        3 * (d[0] + d[3] + d[6])
        + 7 * (d[1] + d[4] + d[7])
        + 1 * (d[2] + d[5] + d[8])
    )
    return checksum % 10 == 0


def iban_valid(s: str) -> bool:
    s = "".join(s.split()).upper()
    if len(s) < 15 or len(s) > 34:
        return False
    if not (s[:2].isalpha() and s[2:4].isdigit()):
        return False
    rearranged = s[4:] + s[:4]
    expanded: list[str] = []
    for ch in rearranged:
        if ch.isdigit():
            expanded.append(ch)
        elif ch.isalpha():
            expanded.append(str(ord(ch) - 55))
        else:
            return False
    return int("".join(expanded)) % 97 == 1


# ---------------------------------------------------------------------------
# Patterns
# ---------------------------------------------------------------------------

SSN_RE = re.compile(r"\b(?!000|666|9\d\d)\d{3}[-\s]?(?!00)\d{2}[-\s]?(?!0000)\d{4}\b")
EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(
    r"(?<!\d)(?:\+?1[-.\s]?)?\(?([2-9]\d{2})\)?[-.\s]?([2-9]\d{2})[-.\s]?(\d{4})(?!\d)"
)
CC_RE = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")
DATE_RE = re.compile(
    r"\b((?:19|20)\d{2}-(?:0[1-9]|1[0-2])-(?:0[1-9]|[12]\d|3[01])"
    r"|(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)\d{2})\b"
)
IPV4_RE = re.compile(
    r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\b"
)
IPV6_RE = re.compile(
    r"(?<![0-9A-Fa-f:])"
    r"(?:"
    r"(?:[0-9A-Fa-f]{1,4}:){7}[0-9A-Fa-f]{1,4}"
    r"|(?:[0-9A-Fa-f]{1,4}:){1,7}:"
    r"|(?:[0-9A-Fa-f]{1,4}:){1,6}:[0-9A-Fa-f]{1,4}"
    r"|(?:[0-9A-Fa-f]{1,4}:){1,5}(?::[0-9A-Fa-f]{1,4}){1,2}"
    r"|(?:[0-9A-Fa-f]{1,4}:){1,4}(?::[0-9A-Fa-f]{1,4}){1,3}"
    r"|(?:[0-9A-Fa-f]{1,4}:){1,3}(?::[0-9A-Fa-f]{1,4}){1,4}"
    r"|(?:[0-9A-Fa-f]{1,4}:){1,2}(?::[0-9A-Fa-f]{1,4}){1,5}"
    r"|[0-9A-Fa-f]{1,4}:(?:(?::[0-9A-Fa-f]{1,4}){1,6})"
    r"|:(?:(?::[0-9A-Fa-f]{1,4}){1,7}|:)"
    r")"
    r"(?![0-9A-Fa-f:])"
)
ROUTING_RE = re.compile(r"\b\d{9}\b")
PHONE_INTL_RE = re.compile(
    r"(?<![\d+])\+(?!1[\s.-]?\d)(?:\d[\s.-]?){7,15}\d(?!\d)"
)
IBAN_RE = re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b")
PASSPORT_INLINE_RE = re.compile(
    r"(?i)\bpassport(?:\s*(?:no|number|#))?[^\w]{0,10}([A-Z0-9]\d{8}|\d{9})\b"
)
DL_INLINE_RE = re.compile(
    r"(?i)\b(?:driver'?s?\s*licen[sc]e|dl\s*(?:no|number|#)?)[^\w]{0,10}([A-Z0-9]{4,12})\b"
)
ICD10_INLINE_RE = re.compile(
    r"(?i)\b(?:icd-?10|diagnosis|dx)[^\w]{0,10}([A-Z]\d{2}(?:\.[A-Z0-9]{1,4})?)\b"
)
AWS_ACCESS_KEY_RE = re.compile(r"\b(?:AKIA|ASIA|AGPA|AIDA|AROA|AIPA|ANPA|ANVA)[0-9A-Z]{16}\b")
AWS_SECRET_INLINE_RE = re.compile(
    r"(?i)aws[_\-\s]?secret(?:[_\-\s]?access)?(?:[_\-\s]?key)?"
    r"[\s:=\"']{0,10}([A-Za-z0-9/+=]{40})\b"
)
GITHUB_TOKEN_RE = re.compile(r"\bgh[pousr]_[A-Za-z0-9]{36,}\b")
SLACK_TOKEN_RE = re.compile(r"\bxox[abprs]-[A-Za-z0-9-]{10,}\b")
STRIPE_KEY_RE = re.compile(r"\b(?:sk|rk|pk)_(?:live|test)_[A-Za-z0-9]{20,}\b")
GOOGLE_API_KEY_RE = re.compile(r"\bAIza[0-9A-Za-z_\-]{35}\b")
JWT_RE = re.compile(r"\beyJ[A-Za-z0-9_\-]+\.eyJ[A-Za-z0-9_\-]+\.[A-Za-z0-9_\-]+\b")
PEM_PRIVATE_KEY_RE = re.compile(
    r"-----BEGIN (?:RSA |EC |DSA |OPENSSH |PGP |ENCRYPTED )?PRIVATE KEY-----"
)
NAME_INLINE_RE = re.compile(
    r"(?im)\b(?:patient(?:\s*name)?|full[\s_-]?name|first[\s_-]?name|last[\s_-]?name|name)"
    r"\s*[:=]\s*([A-Z][A-Za-z'\-]+(?:\s+[A-Z][A-Za-z'\-]+){1,3})"
)
ADDRESS_INLINE_RE = re.compile(
    r"(?im)\b(?:address|addr|street)\s*[:=]\s*"
    r"(\d+\s+[A-Za-z][A-Za-z0-9'.\-\s]{2,80}"
    r"(?:\s(?:st|street|ave|avenue|blvd|boulevard|rd|road|dr|drive|ln|lane|way|ct|court|pl|place))\b\.?)"
)
DOB_INLINE_RE = re.compile(
    r"(?i)\b(?:dob|date[\s_-]?of[\s_-]?birth|birth[\s_-]?date|birthday)"
    r"\s*[:=]?\s*"
    r"((?:19|20)\d{2}-(?:0?[1-9]|1[0-2])-(?:0?[1-9]|[12]\d|3[01])"
    r"|(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)?\d{2})"
)
MRN_RE = re.compile(r"\b\d{6,10}\b")
MRN_INLINE_RE = re.compile(
    r"(?i)\b(?:mrn|medical[\s_-]?record(?:\s*(?:no|number|#))?|patient[\s_-]?id)"
    r"[^\w]{0,20}(\d{6,10})\b"
)

VALUE_DETECTORS: list[tuple[str, "re.Pattern[str]", Optional[Callable[[str], bool]], str, str]] = [
    ("AWS_ACCESS_KEY_ID", AWS_ACCESS_KEY_RE, None, "regex", "high"),
    ("GITHUB_TOKEN", GITHUB_TOKEN_RE, None, "regex", "high"),
    ("SLACK_TOKEN", SLACK_TOKEN_RE, None, "regex", "high"),
    ("STRIPE_KEY", STRIPE_KEY_RE, None, "regex", "high"),
    ("GOOGLE_API_KEY", GOOGLE_API_KEY_RE, None, "regex", "high"),
    ("JWT", JWT_RE, None, "regex", "high"),
    ("PRIVATE_KEY", PEM_PRIVATE_KEY_RE, None, "regex", "high"),
    ("SSN", SSN_RE, ssn_valid, "regex+validity", "high"),
    ("CREDIT_CARD", CC_RE, luhn_valid, "regex+luhn", "high"),
    ("EMAIL", EMAIL_RE, None, "regex", "high"),
    ("PHONE_US", PHONE_RE, None, "regex", "high"),
    ("PHONE_INTL", PHONE_INTL_RE, None, "regex", "medium"),
    ("IBAN", IBAN_RE, iban_valid, "regex+checksum", "high"),
    ("DATE", DATE_RE, None, "regex", "medium"),
    ("IP_ADDRESS", IPV4_RE, None, "regex", "medium"),
    ("IPV6_ADDRESS", IPV6_RE, None, "regex", "medium"),
    ("ROUTING_NUMBER", ROUTING_RE, routing_valid, "regex+checksum", "medium"),
]

INLINE_DETECTORS: list[tuple[str, "re.Pattern[str]", Optional[Callable[[str], bool]], str, str]] = [
    ("PASSPORT", PASSPORT_INLINE_RE, None, "regex+inline_keyword", "high"),
    ("DRIVERS_LICENSE", DL_INLINE_RE, None, "regex+inline_keyword", "medium"),
    ("ICD10", ICD10_INLINE_RE, None, "regex+inline_keyword", "medium"),
    ("AWS_SECRET_ACCESS_KEY", AWS_SECRET_INLINE_RE, None, "regex+inline_keyword", "high"),
]

NAME_ADDRESS_DOB_INLINE: list[tuple[str, "re.Pattern[str]", str, str]] = [
    ("NAME", NAME_INLINE_RE, "regex+inline_keyword", "medium"),
    ("ADDRESS", ADDRESS_INLINE_RE, "regex+inline_keyword", "medium"),
    ("DOB", DOB_INLINE_RE, "regex+inline_keyword", "high"),
]


# ---------------------------------------------------------------------------
# Header rules
# ---------------------------------------------------------------------------

HEADER_RULES: list[tuple["re.Pattern[str]", str]] = [
    (re.compile(r"(^|_)ssn($|_)"), "SSN"),
    (re.compile(r"social.*sec"), "SSN"),
    (re.compile(r"(^|_)mrn($|_)"), "MRN"),
    (re.compile(r"medical.*record"), "MRN"),
    (re.compile(r"patient.*id"), "PATIENT_ID"),
    (re.compile(r"(^|_)npi($|_)"), "NPI"),
    (re.compile(r"(^|_)dob($|_)"), "DOB"),
    (re.compile(r"date.*of.*birth|birth.*date|birthdate"), "DOB"),
    (re.compile(r"first.*name|fname|given.*name"), "NAME"),
    (re.compile(r"last.*name|lname|surname|family.*name"), "NAME"),
    (re.compile(r"patient.*name|full.*name|(^|_)name($|_)"), "NAME"),
    (re.compile(r"middle.*name|middle.*initial"), "NAME"),
    (re.compile(r"street|address|addr"), "ADDRESS"),
    (re.compile(r"(^|_)city($|_)"), "ADDRESS"),
    (re.compile(r"(^|_)state($|_)"), "ADDRESS"),
    (re.compile(r"(^|_)zip($|_)|postal.*code"), "ZIP_CODE"),
    (re.compile(r"(^|_)phone($|_)|phone.*number|mobile|cell"), "PHONE_US"),
    (re.compile(r"(^|_)email($|_)|e.?mail"), "EMAIL"),
    (re.compile(r"diagnosis|icd.?10|icd.?9"), "DIAGNOSIS"),
    (re.compile(r"insurance|policy.*num"), "INSURANCE"),
    (re.compile(r"credit.*card|(^|_)cc.*num|card.*number"), "CREDIT_CARD"),
    (re.compile(r"(^|_)iban($|_)|bank.*account|account.*number"), "IBAN"),
    (re.compile(r"(^|_)swift($|_)|bic.*code"), "SWIFT"),
    (re.compile(r"passport"), "PASSPORT"),
    (re.compile(r"driver.?s?.?licen[sc]e|(^|_)dl.?(num|number|no)?($|_)"), "DRIVERS_LICENSE"),
    (re.compile(r"(^|_)api.?key($|_)|secret.?key|access.?token|auth.?token"), "SECRET"),
    (re.compile(r"(^|_)password($|_)|passwd|pwd"), "PASSWORD"),
    (re.compile(r"routing.?(num|number|no)"), "ROUTING_NUMBER"),
]

_NON_PAN_HEADERS = re.compile(
    r"(?:^|_)(?:tx|txn|transaction|order|invoice|receipt|shipment|tracking|"
    r"barcode|upc|ean|isbn|sku|batch|lot|session|request|trace|correlation"
    r")(?:_id|_num(?:ber)?|_no)?(?:$|_)|tracking_number"
)


def _normalize_header(header: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", header.strip().lower()).strip("_")


def entity_for_header(header: Optional[str]) -> Optional[str]:
    if not header:
        return None
    norm = _normalize_header(header)
    for pat, entity in HEADER_RULES:
        if pat.search(norm):
            return entity
    return None


def suppresses_credit_card(header: Optional[str]) -> bool:
    if not header:
        return False
    return bool(_NON_PAN_HEADERS.search(_normalize_header(header)))


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------

def _mask(s: str) -> str:
    s = s.strip()
    if len(s) <= 4:
        return "*" * len(s)
    return "*" * (len(s) - 4) + s[-4:]


def _snippet(match: str, mask: bool) -> str:
    match = match.strip()
    return _mask(match) if mask else match


def detect(chunk: Chunk, mask: bool = True) -> list[Finding]:
    findings: list[Finding] = []
    text = chunk.text or ""
    seen: set[tuple[int, int]] = set()
    header_entity = entity_for_header(chunk.column_header)

    if header_entity and text.strip():
        findings.append(Finding(
            file_path=chunk.file_path, file_type=chunk.file_type,
            location=chunk.location, entity_type=f"{header_entity}_FROM_HEADER",
            detection_rule="header_match", confidence="high",
            match_snippet=_mask(text) if mask else text,
            column_header=chunk.column_header,
        ))

    # Inline keyword detectors (more specific) before generic value detectors.
    for entity, regex, validator, rule, conf in INLINE_DETECTORS:
        for m in regex.finditer(text):
            sp = m.span(1)
            if sp in seen:
                continue
            v = m.group(1)
            if validator and not validator(v):
                continue
            seen.add(sp)
            findings.append(Finding(
                file_path=chunk.file_path, file_type=chunk.file_type,
                location=chunk.location, entity_type=entity,
                detection_rule=rule, confidence=conf,
                match_snippet=_snippet(v, mask),
                column_header=chunk.column_header,
            ))

    for entity, regex, rule, conf in NAME_ADDRESS_DOB_INLINE:
        for m in regex.finditer(text):
            sp = m.span(1)
            if sp in seen:
                continue
            seen.add(sp)
            findings.append(Finding(
                file_path=chunk.file_path, file_type=chunk.file_type,
                location=chunk.location, entity_type=entity,
                detection_rule=rule, confidence=conf,
                match_snippet=_snippet(m.group(1), mask),
                column_header=chunk.column_header,
            ))

    for m in MRN_INLINE_RE.finditer(text):
        sp = m.span(1)
        if sp in seen:
            continue
        seen.add(sp)
        findings.append(Finding(
            file_path=chunk.file_path, file_type=chunk.file_type,
            location=chunk.location, entity_type="MRN",
            detection_rule="regex+inline_keyword", confidence="high",
            match_snippet=_snippet(m.group(1), mask),
            column_header=chunk.column_header,
        ))

    if header_entity in {"MRN", "PATIENT_ID"}:
        for m in MRN_RE.finditer(text):
            sp = m.span()
            if sp in seen:
                continue
            seen.add(sp)
            findings.append(Finding(
                file_path=chunk.file_path, file_type=chunk.file_type,
                location=chunk.location, entity_type="MRN",
                detection_rule="regex+header_context", confidence="high",
                match_snippet=_snippet(m.group(0), mask),
                column_header=chunk.column_header,
            ))

    cc_suppressed = suppresses_credit_card(chunk.column_header)
    for entity, regex, validator, rule, conf in VALUE_DETECTORS:
        if entity == "CREDIT_CARD" and cc_suppressed:
            continue
        for m in regex.finditer(text):
            sp = m.span()
            if sp in seen:
                continue
            matched = m.group(0)
            if validator and not validator(matched):
                continue
            seen.add(sp)
            findings.append(Finding(
                file_path=chunk.file_path, file_type=chunk.file_type,
                location=chunk.location, entity_type=entity,
                detection_rule=rule, confidence=conf,
                match_snippet=_snippet(matched, mask),
                column_header=chunk.column_header,
            ))

    return findings


# ---------------------------------------------------------------------------
# Readers — CSV and text are stdlib-only. Others lazy-import their dep.
# ---------------------------------------------------------------------------

def read_csv(path: Path) -> Iterator[Chunk]:
    with path.open("r", encoding="utf-8", errors="replace", newline="") as fh:
        reader = _csv.reader(fh)
        try:
            headers = next(reader)
        except StopIteration:
            return
        for row_num, row in enumerate(reader, start=2):
            for col_idx, cell in enumerate(row):
                if cell is None or cell == "":
                    continue
                header = headers[col_idx] if col_idx < len(headers) else None
                yield Chunk(
                    file_path=str(path), file_type="csv",
                    location=f"row={row_num} col={header or col_idx + 1}",
                    text=str(cell), column_header=header,
                )


def read_text(path: Path) -> Iterator[Chunk]:
    file_type = path.suffix.lstrip(".").lower() or "text"
    with path.open("r", encoding="utf-8", errors="replace") as fh:
        for line_num, line in enumerate(fh, start=1):
            line = line.rstrip("\n")
            if not line:
                continue
            yield Chunk(
                file_path=str(path), file_type=file_type,
                location=f"line={line_num}", text=line,
            )


def read_excel(path: Path) -> Iterator[Chunk]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise RuntimeError("openpyxl is required for .xlsx files: pip install openpyxl")
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers: list[Optional[str]] = []
            for row_num, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if row_num == 1:
                    headers = [str(c) if c is not None else None for c in row]
                    continue
                for col_idx, cell in enumerate(row):
                    if cell is None or cell == "":
                        continue
                    header = headers[col_idx] if col_idx < len(headers) else None
                    yield Chunk(
                        file_path=str(path), file_type="xlsx",
                        location=f"sheet={sheet_name} row={row_num} col={header or col_idx + 1}",
                        text=str(cell), column_header=header,
                    )
    finally:
        wb.close()


# Minimal SQL reader: regex-based string and comment extraction. Doesn't need
# sqlparse so the single-file version stays stdlib-only for SQL.
_SQL_STRING = re.compile(r"'((?:[^']|'')*)'")
_SQL_LINE_COMMENT = re.compile(r"--[^\n]*")
_SQL_BLOCK_COMMENT = re.compile(r"/\*.*?\*/", re.DOTALL)


def read_sql(path: Path) -> Iterator[Chunk]:
    source = path.read_text(encoding="utf-8", errors="replace")

    def line_of(offset: int) -> int:
        return source.count("\n", 0, offset) + 1

    for m in _SQL_STRING.finditer(source):
        inner = m.group(1).replace("''", "'")
        if inner:
            yield Chunk(
                file_path=str(path), file_type="sql",
                location=f"line={line_of(m.start())}", text=inner,
            )
    for pat, label in ((_SQL_LINE_COMMENT, "comment"), (_SQL_BLOCK_COMMENT, "comment")):
        for m in pat.finditer(source):
            text = m.group(0)
            if text.strip():
                yield Chunk(
                    file_path=str(path), file_type="sql",
                    location=f"line={line_of(m.start())} ({label})",
                    text=text,
                )


def _walk_json(value: Any, path: str, header: Optional[str]) -> Iterator[tuple[str, str, Optional[str]]]:
    if value is None:
        return
    if isinstance(value, dict):
        for k, v in value.items():
            key = str(k)
            yield from _walk_json(v, f"{path}.{key}" if path else key, header=key)
    elif isinstance(value, list):
        for i, item in enumerate(value):
            yield from _walk_json(item, f"{path}[{i}]", header=header)
    else:
        s = str(value)
        if s == "":
            return
        yield (path or "$", s, header)


def read_json(path: Path) -> Iterator[Chunk]:
    suffix = path.suffix.lower()
    file_type = suffix.lstrip(".") or "json"

    if suffix == ".ndjson":
        with path.open("r", encoding="utf-8", errors="replace") as fh:
            for line_num, line in enumerate(fh, start=1):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except json.JSONDecodeError:
                    continue
                for loc, text, header in _walk_json(obj, "", None):
                    yield Chunk(
                        file_path=str(path), file_type=file_type,
                        location=f"line={line_num} {loc}",
                        text=text, column_header=header,
                    )
        return

    raw = path.read_text(encoding="utf-8", errors="replace")
    if suffix in {".yaml", ".yml"}:
        try:
            import yaml  # type: ignore
        except ImportError:
            raise RuntimeError("PyYAML is required for .yaml/.yml files: pip install PyYAML")
        loaded = yaml.safe_load(raw)
    else:
        loaded = json.loads(raw)

    for loc, text, header in _walk_json(loaded, "", None):
        yield Chunk(
            file_path=str(path), file_type=file_type,
            location=loc, text=text, column_header=header,
        )


_PDFMINER_LOG_MUTED = False


def _mute_pdfminer_logging() -> None:
    """Silence pdfminer's chatty WARNING-level diagnostics ('Could not get
    FontBBox...', 'metadata field indicating that it should not allow text
    extraction...'). Users who want them can set PDFMINER_LOG=1."""
    global _PDFMINER_LOG_MUTED
    if _PDFMINER_LOG_MUTED:
        return
    import logging
    import os
    _PDFMINER_LOG_MUTED = True
    if os.environ.get("PDFMINER_LOG"):
        return
    for name in ("pdfminer", "pdfminer.pdffont", "pdfminer.pdfinterp",
                 "pdfminer.pdfpage", "pdfminer.pdfparser", "pdfminer.cmapdb",
                 "pdfminer.layout", "pdfminer.converter"):
        logging.getLogger(name).setLevel(logging.ERROR)


def read_pdf(path: Path) -> Iterator[Chunk]:
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer
    except ImportError:
        raise RuntimeError("pdfminer.six is required for .pdf files: pip install pdfminer.six")
    _mute_pdfminer_logging()
    for page_num, page in enumerate(extract_pages(str(path)), start=1):
        parts: list[str] = []
        for el in page:
            if isinstance(el, LTTextContainer):
                parts.append(el.get_text())
        text = "".join(parts).strip()
        if text:
            yield Chunk(
                file_path=str(path), file_type="pdf",
                location=f"page={page_num}", text=text,
            )


def read_docx(path: Path) -> Iterator[Chunk]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx is required for .docx files: pip install python-docx")
    doc = Document(str(path))
    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            yield Chunk(
                file_path=str(path), file_type="docx",
                location=f"paragraph={i}", text=text,
            )
    for tbl_idx, table in enumerate(doc.tables, start=1):
        rows = list(table.rows)
        if not rows:
            continue
        header_cells = [c.text.strip() for c in rows[0].cells]
        for row_num, row in enumerate(rows[1:], start=2):
            for col_idx, cell in enumerate(row.cells):
                value = cell.text.strip()
                if not value:
                    continue
                header = header_cells[col_idx] if col_idx < len(header_cells) else None
                yield Chunk(
                    file_path=str(path), file_type="docx",
                    location=f"table={tbl_idx} row={row_num} col={header or col_idx + 1}",
                    text=value, column_header=header,
                )


READERS: dict[str, Callable[[Path], Iterator[Chunk]]] = {
    ".csv": read_csv,
    ".xlsx": read_excel,
    ".sql": read_sql,
    ".txt": read_text,
    ".log": read_text,
    ".md": read_text,
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".json": read_json,
    ".ndjson": read_json,
    ".yaml": read_json,
    ".yml": read_json,
}

DEFAULT_EXTS = tuple(READERS.keys())


# ---------------------------------------------------------------------------
# Scanner
# ---------------------------------------------------------------------------

def _iter_files(
    roots: Iterable[Path],
    include_ext: Iterable[str],
    excludes: Iterable[str],
    max_size_bytes: int,
    verbose: bool,
) -> Iterator[Path]:
    exts = {e.lower() for e in include_ext}
    excludes = list(excludes)
    for root in roots:
        if root.is_file():
            if root.suffix.lower() in exts:
                yield root
            continue
        for path in root.rglob("*"):
            if not path.is_file():
                continue
            if path.suffix.lower() not in exts:
                continue
            rel = str(path)
            if any(fnmatch.fnmatch(rel, pat) for pat in excludes):
                continue
            try:
                if path.stat().st_size > max_size_bytes:
                    if verbose:
                        print(f"skip (too large): {path}", file=sys.stderr)
                    continue
            except OSError:
                continue
            yield path


def scan(
    roots: Iterable[Path],
    include_ext: Iterable[str] = DEFAULT_EXTS,
    excludes: Iterable[str] = (),
    max_size_mb: int = 200,
    mask: bool = True,
    verbose: bool = False,
    errors_out: Optional[list[str]] = None,
) -> Iterator[Finding]:
    max_bytes = max_size_mb * 1024 * 1024
    for path in _iter_files(roots, include_ext, excludes, max_bytes, verbose):
        reader = READERS.get(path.suffix.lower())
        if reader is None:
            continue
        if verbose:
            print(f"scanning: {path}", file=sys.stderr)
        try:
            for chunk in reader(path):
                yield from detect(chunk, mask=mask)
        except Exception as e:
            msg = f"error scanning {path}: {type(e).__name__}: {e}"
            print(msg, file=sys.stderr)
            if errors_out is not None:
                errors_out.append(msg)


# ---------------------------------------------------------------------------
# Git-history scanner
# ---------------------------------------------------------------------------

def _git(repo: Path, *args: str) -> bytes:
    r = subprocess.run(["git", "-C", str(repo), *args], capture_output=True)
    if r.returncode != 0:
        raise RuntimeError(f"git {' '.join(args)} failed: {r.stderr.decode(errors='replace')}")
    return r.stdout


def is_git_repo(path: Path) -> bool:
    try:
        subprocess.run(["git", "-C", str(path), "rev-parse", "--git-dir"],
                       capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


def iter_deletions(repo: Path, include_ext: tuple[str, ...]) -> Iterator[tuple[str, str]]:
    out = _git(repo, "log", "--all", "--diff-filter=D",
               "--name-only", "--pretty=format:%H").decode(errors="replace")
    exts = {e.lower() for e in include_ext}
    current: Optional[str] = None
    seen: set[tuple[str, str]] = set()
    for raw in out.splitlines():
        line = raw.strip()
        if not line:
            continue
        if len(line) == 40 and all(c in "0123456789abcdef" for c in line):
            current = line
            continue
        if current is None:
            continue
        ext = "." + line.rsplit(".", 1)[-1].lower() if "." in line else ""
        if ext not in exts:
            continue
        key = (current, line)
        if key in seen:
            continue
        seen.add(key)
        yield key


def scan_git_deleted(
    repo: Path,
    include_ext: tuple[str, ...],
    max_size_mb: int = 200,
    mask: bool = True,
    verbose: bool = False,
    max_commits: Optional[int] = None,
    errors_out: Optional[list[str]] = None,
) -> Iterator[Finding]:
    if not is_git_repo(repo):
        raise RuntimeError(f"not a git repository: {repo}")
    max_bytes = max_size_mb * 1024 * 1024
    processed = 0
    for commit, path in iter_deletions(repo, include_ext):
        if max_commits is not None and processed >= max_commits:
            break
        processed += 1
        r = subprocess.run(["git", "-C", str(repo), "show", f"{commit}^:{path}"],
                           capture_output=True)
        if r.returncode != 0:
            if verbose:
                print(f"skip (no parent or missing blob): {commit[:10]} {path}", file=sys.stderr)
            continue
        blob = r.stdout
        if len(blob) > max_bytes:
            if verbose:
                print(f"skip (too large): {commit[:10]} {path}", file=sys.stderr)
            continue
        suffix = Path(path).suffix.lower()
        reader = READERS.get(suffix)
        if reader is None:
            continue
        if verbose:
            print(f"scanning: {commit[:10]} {path}", file=sys.stderr)
        with tempfile.NamedTemporaryFile(prefix="untidy-git-", suffix=suffix, delete=False) as tf:
            tf.write(blob)
            tmp = Path(tf.name)
        try:
            for chunk in reader(tmp):
                for f in detect(chunk, mask=mask):
                    yield Finding(
                        file_path=f"git:{path}", file_type=f.file_type,
                        location=f"commit={commit[:10]} {f.location}",
                        entity_type=f.entity_type, detection_rule=f.detection_rule,
                        confidence=f.confidence, match_snippet=f.match_snippet,
                        column_header=f.column_header,
                    )
        except Exception as e:
            msg = f"error scanning {commit[:10]} {path}: {type(e).__name__}: {e}"
            print(msg, file=sys.stderr)
            if errors_out is not None:
                errors_out.append(msg)
        finally:
            try:
                tmp.unlink()
            except OSError:
                pass


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------

REPORT_COLUMNS = [
    "file_path", "file_type", "location", "entity_type",
    "detection_rule", "confidence", "match_snippet", "column_header",
]


def write_csv(
    findings: Iterable[Finding], output_path: Path, mask: bool = True
) -> int:
    """Write findings to CSV. mask=True applies _mask() to match_snippet at
    write time, so a single scan can emit both masked and unmasked reports."""
    count = 0
    with output_path.open("w", encoding="utf-8", newline="") as fh:
        w = _csv.DictWriter(fh, fieldnames=REPORT_COLUMNS)
        w.writeheader()
        for f in findings:
            snippet = _mask(f.match_snippet) if mask else f.match_snippet
            w.writerow({
                "file_path": f.file_path, "file_type": f.file_type,
                "location": f.location, "entity_type": f.entity_type,
                "detection_rule": f.detection_rule, "confidence": f.confidence,
                "match_snippet": snippet,
                "column_header": f.column_header or "",
            })
            count += 1
    return count


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

_CONFIDENCE_RANK = {"low": 0, "medium": 1, "high": 2}


def _parse_exts(raw: str) -> list[str]:
    return [
        e.strip() if e.strip().startswith(".") else f".{e.strip()}"
        for e in raw.split(",") if e.strip()
    ]


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="untidy_solo",
        description="Single-file offline scanner for sensitive data.",
    )
    p.add_argument("--version", action="version", version=f"untidy_solo {__version__}")
    sub = p.add_subparsers(dest="cmd", required=True)

    s = sub.add_parser("scan", help="Scan paths for sensitive data")
    s.add_argument("paths", nargs="+", type=Path)
    s.add_argument("--output", type=Path, default=Path("untidy-findings.csv"),
                   help="Masked CSV report path")
    s.add_argument("--unmasked-output", type=Path, default=None,
                   help="Optional second CSV with raw match values for triage. "
                        "Treat this file as PHI/PII.")
    s.add_argument("--exclude", action="append", default=[], metavar="GLOB")
    s.add_argument("--include-ext", default=",".join(DEFAULT_EXTS))
    s.add_argument("--max-size-mb", type=int, default=200)
    s.add_argument("--min-confidence", choices=["low", "medium", "high"], default="low")
    s.add_argument("--no-mask", action="store_true",
                   help="Emit raw matches in --output instead of masked values")
    s.add_argument("--strict", action="store_true",
                   help="Exit non-zero if any file failed to read")
    s.add_argument("--verbose", action="store_true")

    g = sub.add_parser("scan-git", help="Scan git history for deleted sensitive files")
    g.add_argument("repo", type=Path)
    g.add_argument("--output", type=Path, default=Path("untidy-findings.csv"))
    g.add_argument("--unmasked-output", type=Path, default=None,
                   help="Optional second CSV with raw match values for triage.")
    g.add_argument("--include-ext", default=",".join(DEFAULT_EXTS))
    g.add_argument("--max-size-mb", type=int, default=200)
    g.add_argument("--max-commits", type=int, default=None)
    g.add_argument("--min-confidence", choices=["low", "medium", "high"], default="low")
    g.add_argument("--no-mask", action="store_true")
    g.add_argument("--strict", action="store_true")
    g.add_argument("--verbose", action="store_true")

    return p


def main(argv: Optional[list[str]] = None) -> int:
    args = _build_parser().parse_args(argv)
    threshold = _CONFIDENCE_RANK[args.min_confidence]
    exts = _parse_exts(args.include_ext)
    errors: list[str] = []

    if args.cmd == "scan":
        for p in args.paths:
            if not p.exists():
                print(f"error: path does not exist: {p}", file=sys.stderr)
                return 2
        source = scan(
            roots=args.paths, include_ext=exts, excludes=args.exclude,
            max_size_mb=args.max_size_mb, mask=False,
            verbose=args.verbose, errors_out=errors,
        )
    elif args.cmd == "scan-git":
        if not args.repo.exists():
            print(f"error: path does not exist: {args.repo}", file=sys.stderr)
            return 2
        if not is_git_repo(args.repo):
            print(f"error: not a git repository: {args.repo}", file=sys.stderr)
            return 2
        source = scan_git_deleted(
            repo=args.repo, include_ext=tuple(exts),
            max_size_mb=args.max_size_mb, mask=False,
            verbose=args.verbose, max_commits=args.max_commits,
            errors_out=errors,
        )
    else:
        return 2

    findings = [f for f in source if _CONFIDENCE_RANK[f.confidence] >= threshold]

    primary_masked = not args.no_mask
    count = write_csv(findings, args.output, mask=primary_masked)
    label = "masked" if primary_masked else "unmasked"
    print(f"wrote {count} finding(s) to {args.output} ({label})", file=sys.stderr)

    if args.unmasked_output is not None:
        if args.unmasked_output == args.output:
            print("error: --unmasked-output must differ from --output",
                  file=sys.stderr)
            return 2
        write_csv(findings, args.unmasked_output, mask=False)
        print(
            f"wrote {count} finding(s) to {args.unmasked_output} (unmasked — contains PHI/PII)",
            file=sys.stderr,
        )

    if errors:
        print(f"{len(errors)} file(s) had read errors", file=sys.stderr)
        if args.strict:
            return 2
    return 1 if count else 0


if __name__ == "__main__":
    sys.exit(main())
